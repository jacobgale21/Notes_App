import io
from pypdf import PdfReader
from fastapi import HTTPException, UploadFile, File
from docx import Document

def read_pdf_file(pdf_file: UploadFile) -> str:
    if pdf_file.content_type not in ("application/pdf", "application/x-pdf"):
        raise HTTPException(status_code=400, detail="File must be a PDF")
    
    data = pdf_file.file.read() 
    if not data:
        raise HTTPException(status_code=400, detail="Empty file")

    try:
        reader = PdfReader(io.BytesIO(data))
        pages = [page.extract_text() or "" for page in reader.pages]
    except Exception as e:
        raise HTTPException(status_code=400, detail="Could not read PDF") from e

    text = "\n".join(pages).strip()
    if not text:
        raise HTTPException(
            status_code=400,
            detail="No text in this PDF (it may be scanned images)",
        )
    return text

def read_word_file(word_file: UploadFile) -> str:
    contents = word_file.file.read()
    if not contents:
        raise HTTPException(status_code=400, detail="Empty file")
    try:
        doc = Document(io.BytesIO(contents))
        paragraphs = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                paragraphs.append(" ".join(cell.text for cell in row.cells))

        text = "\n".join(paragraphs).strip()
        if not text:
            raise HTTPException(status_code=400, detail="No text in this Word document")
        return text
    except Exception as e:
        raise HTTPException(status_code=400, detail="Could not read Word document") from e
